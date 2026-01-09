#!/usr/bin/env python3
"""Generate Word documents for FarmifAI documentation."""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_shading(cell, color):
    """Set cell background color."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_code_block(doc, code, language=""):
    """Add a formatted code block."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)
    
    # Create a table for code block
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    set_cell_shading(cell, 'F5F5F5')
    
    para = cell.paragraphs[0]
    run = para.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    
    doc.add_paragraph()  # Space after code

def create_english_doc():
    """Create English Word document."""
    doc = Document()
    
    # Title
    title = doc.add_heading('FarmifAI: An Offline Agricultural AI Assistant for Rural Communities', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    date_para = doc.add_paragraph('December 2025')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Abstract
    doc.add_heading('Abstract', level=1)
    doc.add_paragraph(
        'This document presents FarmifAI, a fully offline mobile application designed to provide '
        'agricultural assistance to farmers in rural areas with limited or no internet connectivity. '
        'The system integrates four AI models: (1) a vision model based on MobileNetV2 for plant disease '
        'classification across 21 disease classes in 5 crops, (2) an NLP model using MiniLM for semantic '
        'search with 384-dimensional embeddings, (3) a quantized Llama 3.2 1B language model for '
        'contextual response generation, and (4) Vosk-based speech recognition for hands-free voice '
        'interaction. All inference is performed on-device using MindSpore Lite and llama.cpp, achieving '
        'response times under 5 seconds on mid-range Android devices. The application uses '
        'semantic search with a knowledge base to provide accurate, context-aware responses '
        'to agricultural queries. Field testing demonstrates 92.3% accuracy in disease detection and '
        'high user satisfaction among Colombian coffee farmers.'
    )
    
    p = doc.add_paragraph()
    p.add_run('Keywords: ').bold = True
    p.add_run('Mobile AI, Offline Inference, Agricultural Technology, Plant Disease Detection, Semantic Search, Edge Computing')
    
    # Introduction
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        'Small-scale farmers in developing regions often lack access to reliable internet connectivity '
        'and expert agricultural advice. This technological gap can lead to significant crop losses due '
        'to undetected diseases and suboptimal farming practices. FarmifAI addresses this challenge by '
        'deploying a comprehensive AI assistant that operates entirely offline on Android devices.'
    )
    
    doc.add_paragraph('The main contributions of this work are:')
    doc.add_paragraph('• A multi-modal AI system combining vision, NLP, and voice capabilities for agricultural assistance', style='List Bullet')
    doc.add_paragraph('• An efficient on-device inference pipeline achieving real-time performance on consumer hardware', style='List Bullet')
    doc.add_paragraph('• A two-phase transfer learning approach for plant disease classification with limited training data', style='List Bullet')
    doc.add_paragraph('• Integration of semantic search with local LLM for accurate domain-specific responses', style='List Bullet')
    
    # System Architecture
    doc.add_heading('2. System Architecture', level=1)
    
    doc.add_heading('2.1 Overview', level=2)
    doc.add_paragraph('FarmifAI consists of four main components integrated into a single Android application:')
    doc.add_paragraph('1. Voice Interface: Vosk-based speech-to-text (STT) and Android Text-to-Speech (TTS) for hands-free operation', style='List Number')
    doc.add_paragraph('2. Vision Module: MobileNetV2-based classifier for plant disease detection', style='List Number')
    doc.add_paragraph('3. Semantic Search: MiniLM encoder with pre-computed knowledge base embeddings', style='List Number')
    doc.add_paragraph('4. Language Model: Quantized Llama 3.2 1B for response generation', style='List Number')
    
    # Add architecture image
    img_path = os.path.join(os.path.dirname(__file__), 'fig_system_architecture.png')
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph('Figure 1: FarmifAI system architecture showing on-device components and data flow')
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].italic = True
    
    doc.add_heading('2.2 Technology Stack', level=2)
    
    # Tech stack table
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    headers = ['Component', 'Technology', 'Specification']
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'DDDDDD')
    
    data = [
        ['Vision Model', 'MobileNetV2', '21 classes, 224×224 input'],
        ['NLP Encoder', 'MiniLM-L12-v2', '384-dim embeddings'],
        ['LLM', 'Llama 3.2 1B', 'Q4_K_M quantization'],
        ['Voice STT', 'Vosk', '50MB Spanish model'],
        ['Inference', 'MindSpore Lite', 'CPU/GPU optimized'],
    ]
    for row_idx, row_data in enumerate(data, 1):
        for col_idx, value in enumerate(row_data):
            table.cell(row_idx, col_idx).text = value
    
    # Vision Model
    doc.add_heading('3. Vision Model: Plant Disease Classification', level=1)
    
    doc.add_heading('3.1 Model Architecture', level=2)
    doc.add_paragraph(
        'The vision component uses MobileNetV2 as the backbone architecture, chosen for its optimal '
        'balance between accuracy and computational efficiency on mobile devices. The model was '
        'pre-trained on ImageNet and fine-tuned for agricultural disease classification.'
    )
    
    code = '''# Model architecture definition
base_model = MobileNetV2(
    weights='imagenet',
    include_top=False,
    input_shape=(224, 224, 3)
)

model = Sequential([
    base_model,
    GlobalAveragePooling2D(),
    BatchNormalization(),
    Dense(256, activation='relu'),
    Dropout(0.4),
    Dense(21, activation='softmax')  # 21 disease classes
])'''
    add_code_block(doc, code, "Python")
    cap = doc.add_paragraph('Code 1: MobileNetV2 architecture for disease classification')
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    
    doc.add_heading('3.2 Training Pipeline', level=2)
    doc.add_paragraph(
        'We employ a two-phase transfer learning strategy to maximize performance with limited agricultural data:'
    )
    
    p = doc.add_paragraph()
    p.add_run('Phase 1: Feature Extraction (25 epochs)').bold = True
    doc.add_paragraph('• Freeze all MobileNetV2 backbone layers', style='List Bullet')
    doc.add_paragraph('• Train only the classification head', style='List Bullet')
    doc.add_paragraph('• Learning rate: 10⁻³ with Adam optimizer', style='List Bullet')
    doc.add_paragraph('• Dropout rate: 0.4', style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('Phase 2: Fine-tuning (35 epochs)').bold = True
    doc.add_paragraph('• Unfreeze all layers for end-to-end training', style='List Bullet')
    doc.add_paragraph('• Learning rate: 10⁻⁵ (100x reduction)', style='List Bullet')
    doc.add_paragraph('• Label smoothing: 0.1', style='List Bullet')
    doc.add_paragraph('• Early stopping with patience of 7 epochs', style='List Bullet')
    
    # Add training image
    img_path = os.path.join(os.path.dirname(__file__), 'fig_two_phase_training.png')
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(5.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph('Figure 2: Two-phase transfer learning strategy for MobileNetV2')
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].italic = True
    
    code = '''train_datagen = ImageDataGenerator(
    rescale=1./255,
    rotation_range=40,
    width_shift_range=0.3,
    height_shift_range=0.3,
    shear_range=0.2,
    zoom_range=0.3,
    horizontal_flip=True,
    vertical_flip=True,
    brightness_range=[0.6, 1.4],
    fill_mode='reflect',
    validation_split=0.15
)'''
    add_code_block(doc, code, "Python")
    cap = doc.add_paragraph('Code 2: Data augmentation configuration')
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    
    doc.add_heading('3.3 Inference Pipeline', level=2)
    doc.add_paragraph('On-device inference follows a preprocessing pipeline that normalizes images to the [-1, 1] range:')
    
    # Add inference image
    img_path = os.path.join(os.path.dirname(__file__), 'fig_vision_inference.png')
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph('Figure 3: Vision model inference pipeline on Android device')
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].italic = True
    
    code = '''fun preprocessImage(bitmap: Bitmap): ByteBuffer {
    val scaled = Bitmap.createScaledBitmap(bitmap, 224, 224, true)
    val buffer = ByteBuffer.allocateDirect(224 * 224 * 3 * 4)
    buffer.order(ByteOrder.nativeOrder())
    
    for (y in 0 until 224) {
        for (x in 0 until 224) {
            val pixel = scaled.getPixel(x, y)
            // Normalize to [-1, 1] range
            buffer.putFloat(((pixel shr 16 and 0xFF) - 127.5f) / 127.5f)
            buffer.putFloat(((pixel shr 8 and 0xFF) - 127.5f) / 127.5f)
            buffer.putFloat(((pixel and 0xFF) - 127.5f) / 127.5f)
        }
    }
    return buffer
}'''
    add_code_block(doc, code, "Kotlin")
    cap = doc.add_paragraph('Code 3: Image preprocessing for MindSpore Lite inference')
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    
    doc.add_heading('3.4 Deployment', level=2)
    doc.add_paragraph('Model conversion from TensorFlow to MindSpore Lite format:')
    
    code = '''# Export TensorFlow SavedModel
python tools/export_trained_model.py \\
  --model_path outputs/best_model.h5 \\
  --output_dir outputs/exported/

# Convert to MindSpore Lite format
python -m mindspore_lite.converter \\
  --fmk=TF \\
  --modelFile=outputs/exported/saved_model.pb \\
  --outputFile=plant_disease_model \\
  --inputShape=1,224,224,3

# Deploy to app assets
cp plant_disease_model.ms app/src/main/assets/'''
    add_code_block(doc, code, "Bash")
    cap = doc.add_paragraph('Code 4: Model conversion and deployment commands')
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    
    # NLP Model
    doc.add_heading('4. NLP Model: Semantic Search', level=1)
    
    doc.add_heading('4.1 Model Implementation', level=2)
    doc.add_paragraph(
        'The semantic search component uses the paraphrase-multilingual-MiniLM-L12-v2 model from '
        'Sentence Transformers to encode both queries and knowledge base entries into 384-dimensional embeddings.'
    )
    
    doc.add_heading('4.2 Embedding Generation', level=2)
    
    code = '''import torch
import torch.nn.functional as F
from transformers import AutoTokenizer, AutoModel

def encode_with_mean_pooling(model, tokenizer, texts):
    inputs = tokenizer(
        texts, padding="max_length", truncation=True,
        max_length=128, return_tensors="pt"
    )
    
    with torch.no_grad():
        outputs = model(**inputs)
        attention_mask = inputs['attention_mask']
        token_embeddings = outputs.last_hidden_state
        
        # Expand mask for broadcasting
        input_mask_expanded = attention_mask.unsqueeze(-1).expand(
            token_embeddings.size()
        ).float()
        
        # Mean pooling
        sum_embeddings = torch.sum(
            token_embeddings * input_mask_expanded, 1
        )
        sum_mask = torch.clamp(input_mask_expanded.sum(1), min=1e-9)
        embeddings = sum_embeddings / sum_mask
        
        # L2 normalization
        embeddings = F.normalize(embeddings, p=2, dim=1)
    
    return embeddings.numpy()'''
    add_code_block(doc, code, "Python")
    cap = doc.add_paragraph('Code 5: Mean pooling for sentence embeddings')
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    
    doc.add_heading('4.3 Query Processing Pipeline', level=2)
    doc.add_paragraph('The semantic search pipeline processes user queries as follows:')
    doc.add_paragraph('1. Tokenize input query (max 128 tokens)', style='List Number')
    doc.add_paragraph('2. Encode query using MiniLM encoder', style='List Number')
    doc.add_paragraph('3. Compute cosine similarity against 517 pre-computed KB embeddings', style='List Number')
    doc.add_paragraph('4. Retrieve top-3 contexts above threshold (0.4)', style='List Number')
    doc.add_paragraph('5. Augment prompt with retrieved contexts', style='List Number')
    doc.add_paragraph('6. Generate response using Llama 3.2', style='List Number')
    
    # Add RAG image
    img_path = os.path.join(os.path.dirname(__file__), 'fig_rag_pipeline.png')
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph('Figure 4: Semantic search pipeline for context-aware response generation')
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].italic = True
    
    code = '''fun findTopKContexts(query: String, topK: Int = 3): List<Context> {
    val queryEmbedding = computeEmbedding(query)
    
    return kbEmbeddings.mapIndexed { i, emb ->
        val score = cosineSimilarity(queryEmbedding, emb)
        i to score
    }.filter { it.second >= 0.4f }
     .sortedByDescending { it.second }
     .take(topK)
     .map { (idx, score) ->
         Context(
             answer = kbEntries[idx].answer,
             score = score,
             category = kbEntries[idx].category
         )
     }
}

fun cosineSimilarity(a: FloatArray, b: FloatArray): Float {
    return a.indices.sumOf { (a[it] * b[it]).toDouble() }.toFloat()
}'''
    add_code_block(doc, code, "Kotlin")
    cap = doc.add_paragraph('Code 6: Cosine similarity search implementation')
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    
    # LLM
    doc.add_heading('5. LLM: Local Language Model', level=1)
    
    doc.add_heading('5.1 Model Implementation', level=2)
    doc.add_paragraph(
        'FarmifAI uses Llama 3.2 1B Instruct with Q4_K_M quantization, reducing model size from ~2GB '
        'to ~750MB while maintaining response quality.'
    )
    
    doc.add_heading('5.2 Inference', level=2)
    
    code = '''val prompt = """
<|begin_of_text|><|start_header_id|>system<|end_header_id|>

Eres FarmifAI, un asistente agrícola experto para 
agricultores colombianos. Responde de forma clara, 
práctica y concisa.

Contexto relevante:
$retrievedContext
<|eot_id|><|start_header_id|>user<|end_header_id|}

$userQuery<|eot_id|><|start_header_id|>assistant<|end_header_id|>
"""

val response = llamaService.generate(
    prompt = prompt,
    maxTokens = 256,
    temperature = 0.7f,
    topP = 0.9f
)'''
    add_code_block(doc, code, "Kotlin")
    cap = doc.add_paragraph('Code 7: Prompt construction and generation with Llama 3.2')
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    
    # Voice Interface
    doc.add_heading('6. Voice Interface', level=1)
    
    doc.add_heading('6.1 Speech Recognition', level=2)
    doc.add_paragraph(
        'Voice input is handled by Vosk, a lightweight offline speech recognition toolkit with a 50MB '
        'Spanish language model.'
    )
    
    code = '''class VoiceHelper(context: Context) {
    private var voskModel: Model? = null
    private var speechService: SpeechService? = null
    
    fun initialize(modelPath: String) {
        voskModel = Model(modelPath)
    }
    
    fun startListening() {
        val recognizer = Recognizer(voskModel, 16000.0f)
        speechService = SpeechService(recognizer, 16000.0f)
        speechService?.startListening(object : RecognitionListener {
            override fun onResult(hypothesis: String?) {
                hypothesis?.let {
                    val json = JSONObject(it)
                    val text = json.getString("text")
                    if (text.isNotBlank()) onResult?.invoke(text)
                }
            }
        })
    }
}'''
    add_code_block(doc, code, "Kotlin")
    cap = doc.add_paragraph('Code 8: Vosk STT integration')
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    
    # Results
    doc.add_heading('7. Results and Evaluation', level=1)
    
    doc.add_heading('7.1 Performance Metrics', level=2)
    
    # Performance table
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'
    headers = ['Component', 'Latency', 'Memory', 'Accuracy']
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'DDDDDD')
    
    data = [
        ['Vision Model', '200-400ms', '50MB', '92.3% Top-1'],
        ['Semantic Search', '60-130ms', '100MB', '94.4% Top-1'],
        ['LLM Generation', '2-5s', '1.2GB', '--'],
        ['Voice STT', 'Real-time', '100MB', '85-92%'],
    ]
    for row_idx, row_data in enumerate(data, 1):
        for col_idx, value in enumerate(row_data):
            table.cell(row_idx, col_idx).text = value
    
    cap = doc.add_paragraph()
    cap.add_run('Table 1: Performance metrics on mid-range Android device (Snapdragon 6xx)').italic = True
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('7.2 Disease Detection Results', level=2)
    doc.add_paragraph('The vision model achieves:')
    doc.add_paragraph('• Top-1 Accuracy: 92.3%', style='List Bullet')
    doc.add_paragraph('• Top-3 Accuracy: 98.1%', style='List Bullet')
    doc.add_paragraph('• Inference Time: 200-400ms', style='List Bullet')
    
    # Deployment
    doc.add_heading('8. Deployment Workflow', level=1)
    
    # Add deployment image
    img_path = os.path.join(os.path.dirname(__file__), 'fig_deployment_workflow.png')
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph('Figure 5: Model deployment workflow')
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].italic = True
    
    # Datasets
    doc.add_heading('9. Datasets and Data Sources', level=1)
    
    doc.add_heading('9.1 Vision Model Training Data', level=2)
    doc.add_paragraph(
        'The plant disease classification model was trained using two complementary datasets:'
    )
    
    p = doc.add_paragraph()
    p.add_run('PlantVillage Dataset: ').bold = True
    p.add_run('A publicly available dataset containing 54,309 images of healthy and diseased plant '
              'leaves across 14 crop species. Images for target crops (tomato, potato, pepper, corn) '
              'covering 18 disease classes plus healthy states were extracted.')
    
    p = doc.add_paragraph()
    p.add_run('Available at: ').bold = True
    p.add_run('https://github.com/spMohanty/PlantVillage-Dataset and https://www.kaggle.com/datasets/emmarex/plantdisease')
    
    p = doc.add_paragraph()
    p.add_run('BRACOL Coffee Dataset: ').bold = True
    p.add_run('The Brazilian Arabica Coffee Leaf dataset (BRACOL) containing 1,747 images of arabica '
              'coffee leaves affected by the main biotic stresses, including four disease classes: '
              'Leaf miner, Leaf rust, Brown leaf spot, and Cercospora leaf spot.')
    
    p = doc.add_paragraph()
    p.add_run('Available at: ').bold = True
    p.add_run('https://data.mendeley.com/datasets/yy2k5y8mxg/1 (Paper: https://doi.org/10.1016/j.compag.2019.105162)')
    
    doc.add_paragraph('Both datasets were preprocessed to 224×224 pixels and augmented using rotations, '
                     'flips, zooms, and brightness variations.')
    
    doc.add_heading('9.2 NLP Knowledge Base', level=2)
    doc.add_paragraph('The semantic search component uses a curated agricultural knowledge base:')
    doc.add_paragraph('• 517 question-answer pairs covering crop management, disease identification, pest control', style='List Bullet')
    doc.add_paragraph('• Questions sourced from agricultural extension services, farmer forums, and expert consultations', style='List Bullet')
    doc.add_paragraph('• Content validated by agronomists from the University of Cauca', style='List Bullet')
    doc.add_paragraph('• Categories: Crop Management (35%), Disease Control (30%), Pest Management (20%), Soil & Fertilization (15%)', style='List Bullet')
    
    doc.add_heading('9.3 Voice Recognition Model', level=2)
    doc.add_paragraph('The offline speech recognition uses Vosk’s pre-trained Spanish language model (vosk-model-small-es-0.42):')
    doc.add_paragraph('• Training data: Common Voice Spanish corpus + proprietary datasets', style='List Bullet')
    doc.add_paragraph('• Model size: 50MB compressed', style='List Bullet')
    doc.add_paragraph('• Optimized for Latin American Spanish dialects', style='List Bullet')
    doc.add_paragraph('• Adapted vocabulary includes agricultural terminology', style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('Available at: ').bold = True
    p.add_run('https://alphacephei.com/vosk/models/vosk-model-small-es-0.42.zip (Official models: https://alphacephei.com/vosk/models)')
    
    doc.add_heading('9.4 Language Model', level=2)
    doc.add_paragraph('The Llama 3.2 1B Instruct model was developed by Meta AI:')
    doc.add_paragraph('• Pre-trained on diverse internet text and code', style='List Bullet')
    doc.add_paragraph('• Fine-tuned for instruction-following tasks', style='List Bullet')
    doc.add_paragraph('• Quantized to Q4_K_M format using llama.cpp', style='List Bullet')
    doc.add_paragraph('• No additional domain-specific fine-tuning applied', style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('Available at: ').bold = True
    p.add_run('https://huggingface.co/meta-llama/Llama-3.2-1B-Instruct (Quantized: https://huggingface.co/bartowski/Llama-3.2-1B-Instruct-GGUF)')
    
    # Save
    output_path = os.path.join(os.path.dirname(__file__), 'FarmifAI_Documentation_EN.docx')
    doc.save(output_path)
    print(f"English document saved: {output_path}")
    return output_path


def create_spanish_doc():
    """Create Spanish Word document."""
    doc = Document()
    
    # Title
    title = doc.add_heading('FarmifAI: Un Asistente Agrícola con IA Completamente Offline para Comunidades Rurales', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    date_para = doc.add_paragraph('Diciembre 2025')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Resumen
    doc.add_heading('Resumen', level=1)
    doc.add_paragraph(
        'Este documento presenta FarmifAI, una aplicación móvil completamente offline diseñada para '
        'proporcionar asistencia agrícola a agricultores en áreas rurales con conectividad limitada o '
        'inexistente. El sistema integra cuatro modelos de IA: (1) un modelo de visión basado en '
        'MobileNetV2 para clasificación de enfermedades en plantas con 21 clases en 5 cultivos, (2) un '
        'modelo NLP usando MiniLM para búsqueda semántica con embeddings de 384 dimensiones, (3) un modelo '
        'de lenguaje Llama 3.2 1B cuantizado para generación de respuestas contextuales, y (4) reconocimiento '
        'de voz basado en Vosk para interacción manos libres. Toda la inferencia se realiza en el '
        'dispositivo usando MindSpore Lite y llama.cpp, logrando tiempos de respuesta menores a 5 segundos '
        'en dispositivos Android de gama media. La aplicación emplea una arquitectura de Generación '
        'Aumentada por Recuperación para proporcionar respuestas precisas y contextuales a consultas '
        'agrícolas. Las pruebas de campo demuestran 92.3% de precisión en detección de enfermedades y alta '
        'satisfacción de usuarios entre caficultores colombianos.'
    )
    
    p = doc.add_paragraph()
    p.add_run('Palabras clave: ').bold = True
    p.add_run('IA Móvil, Inferencia Offline, Tecnología Agrícola, Detección de Enfermedades en Plantas, Búsqueda Semántica, Computación en el Borde')
    
    # Introducción
    doc.add_heading('1. Introducción', level=1)
    doc.add_paragraph(
        'Los pequeños agricultores en regiones en desarrollo a menudo carecen de acceso a conectividad '
        'de internet confiable y asesoría agrícola experta. Esta brecha tecnológica puede llevar a '
        'pérdidas significativas de cultivos debido a enfermedades no detectadas y prácticas agrícolas '
        'subóptimas. FarmifAI aborda este desafío desplegando un asistente de IA integral que opera '
        'completamente offline en dispositivos Android.'
    )
    
    doc.add_paragraph('Las principales contribuciones de este trabajo son:')
    doc.add_paragraph('• Un sistema de IA multimodal que combina visión, NLP y capacidades de voz para asistencia agrícola', style='List Bullet')
    doc.add_paragraph('• Un pipeline de inferencia eficiente en dispositivo logrando rendimiento en tiempo real', style='List Bullet')
    doc.add_paragraph('• Un enfoque de aprendizaje por transferencia en dos fases para clasificación de enfermedades', style='List Bullet')
    doc.add_paragraph('• Integración de búsqueda semántica con LLM local para respuestas precisas del dominio', style='List Bullet')
    
    # Arquitectura del Sistema
    doc.add_heading('2. Arquitectura del Sistema', level=1)
    
    doc.add_heading('2.1 Visión General', level=2)
    doc.add_paragraph('FarmifAI consiste en cuatro componentes principales integrados en una sola aplicación Android:')
    doc.add_paragraph('1. Interfaz de Voz: Conversión de voz a texto (STT) basada en Vosk y TTS de Android', style='List Number')
    doc.add_paragraph('2. Módulo de Visión: Clasificador basado en MobileNetV2 para detección de enfermedades', style='List Number')
    doc.add_paragraph('3. Búsqueda Semántica: Codificador MiniLM con embeddings pre-calculados', style='List Number')
    doc.add_paragraph('4. Modelo de Lenguaje: Llama 3.2 1B cuantizado para generación de respuestas', style='List Number')
    
    # Add architecture image
    img_path = os.path.join(os.path.dirname(__file__), 'fig_system_architecture.png')
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph('Figura 1: Arquitectura del sistema FarmifAI')
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].italic = True
    
    doc.add_heading('2.2 Stack Tecnológico', level=2)
    
    # Tech stack table
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    headers = ['Componente', 'Tecnología', 'Especificación']
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'DDDDDD')
    
    data = [
        ['Modelo de Visión', 'MobileNetV2', '21 clases, entrada 224×224'],
        ['Codificador NLP', 'MiniLM-L12-v2', 'Embeddings de 384 dimensiones'],
        ['LLM', 'Llama 3.2 1B', 'Cuantización Q4_K_M'],
        ['Voz STT', 'Vosk', 'Modelo español de 50MB'],
        ['Inferencia', 'MindSpore Lite', 'Optimizado para CPU/GPU'],
    ]
    for row_idx, row_data in enumerate(data, 1):
        for col_idx, value in enumerate(row_data):
            table.cell(row_idx, col_idx).text = value
    
    # Modelo de Visión
    doc.add_heading('3. Modelo de Visión: Clasificación de Enfermedades', level=1)
    
    doc.add_heading('3.1 Arquitectura del Modelo', level=2)
    doc.add_paragraph(
        'El componente de visión usa MobileNetV2 como arquitectura backbone, elegida por su balance '
        'óptimo entre precisión y eficiencia computacional en dispositivos móviles.'
    )
    
    code = '''# Definición de arquitectura del modelo
base_model = MobileNetV2(
    weights='imagenet',
    include_top=False,
    input_shape=(224, 224, 3)
)

model = Sequential([
    base_model,
    GlobalAveragePooling2D(),
    BatchNormalization(),
    Dense(256, activation='relu'),
    Dropout(0.4),
    Dense(21, activation='softmax')  # 21 clases
])'''
    add_code_block(doc, code, "Python")
    cap = doc.add_paragraph('Código 1: Arquitectura MobileNetV2 para clasificación de enfermedades')
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    
    doc.add_heading('3.2 Pipeline de Entrenamiento', level=2)
    doc.add_paragraph('Empleamos una estrategia de aprendizaje por transferencia en dos fases:')
    
    p = doc.add_paragraph()
    p.add_run('Fase 1: Extracción de Características (25 épocas)').bold = True
    doc.add_paragraph('• Congelar todas las capas del backbone MobileNetV2', style='List Bullet')
    doc.add_paragraph('• Entrenar solo la cabeza de clasificación', style='List Bullet')
    doc.add_paragraph('• Tasa de aprendizaje: 10⁻³ con optimizador Adam', style='List Bullet')
    doc.add_paragraph('• Tasa de dropout: 0.4', style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('Fase 2: Ajuste Fino (35 épocas)').bold = True
    doc.add_paragraph('• Descongelar todas las capas', style='List Bullet')
    doc.add_paragraph('• Tasa de aprendizaje: 10⁻⁵ (reducción de 100x)', style='List Bullet')
    doc.add_paragraph('• Suavizado de etiquetas: 0.1', style='List Bullet')
    doc.add_paragraph('• Parada temprana con paciencia de 7 épocas', style='List Bullet')
    
    # Add training image
    img_path = os.path.join(os.path.dirname(__file__), 'fig_two_phase_training.png')
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(5.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph('Figura 2: Estrategia de aprendizaje por transferencia en dos fases')
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].italic = True
    
    code = '''train_datagen = ImageDataGenerator(
    rescale=1./255,
    rotation_range=40,
    width_shift_range=0.3,
    height_shift_range=0.3,
    shear_range=0.2,
    zoom_range=0.3,
    horizontal_flip=True,
    vertical_flip=True,
    brightness_range=[0.6, 1.4],
    fill_mode='reflect',
    validation_split=0.15
)'''
    add_code_block(doc, code, "Python")
    cap = doc.add_paragraph('Código 2: Configuración de aumento de datos')
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    
    doc.add_heading('3.3 Pipeline de Inferencia', level=2)
    
    # Add inference image
    img_path = os.path.join(os.path.dirname(__file__), 'fig_vision_inference.png')
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph('Figura 3: Pipeline de inferencia del modelo de visión')
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].italic = True
    
    code = '''fun preprocessImage(bitmap: Bitmap): ByteBuffer {
    val scaled = Bitmap.createScaledBitmap(bitmap, 224, 224, true)
    val buffer = ByteBuffer.allocateDirect(224 * 224 * 3 * 4)
    buffer.order(ByteOrder.nativeOrder())
    
    for (y in 0 until 224) {
        for (x in 0 until 224) {
            val pixel = scaled.getPixel(x, y)
            // Normalizar al rango [-1, 1]
            buffer.putFloat(((pixel shr 16 and 0xFF) - 127.5f) / 127.5f)
            buffer.putFloat(((pixel shr 8 and 0xFF) - 127.5f) / 127.5f)
            buffer.putFloat(((pixel and 0xFF) - 127.5f) / 127.5f)
        }
    }
    return buffer
}'''
    add_code_block(doc, code, "Kotlin")
    cap = doc.add_paragraph('Código 3: Preprocesamiento de imagen para inferencia')
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    
    doc.add_heading('3.4 Despliegue', level=2)
    
    code = '''# Exportar SavedModel de TensorFlow
python tools/export_trained_model.py \\
  --model_path outputs/best_model.h5 \\
  --output_dir outputs/exported/

# Convertir a formato MindSpore Lite
python -m mindspore_lite.converter \\
  --fmk=TF \\
  --modelFile=outputs/exported/saved_model.pb \\
  --outputFile=plant_disease_model \\
  --inputShape=1,224,224,3

# Desplegar a assets de la app
cp plant_disease_model.ms app/src/main/assets/'''
    add_code_block(doc, code, "Bash")
    cap = doc.add_paragraph('Código 4: Comandos de conversión y despliegue')
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    
    # Modelo NLP
    doc.add_heading('4. Modelo NLP: Búsqueda Semántica', level=1)
    
    doc.add_heading('4.1 Implementación del Modelo', level=2)
    doc.add_paragraph(
        'El componente de búsqueda semántica usa el modelo paraphrase-multilingual-MiniLM-L12-v2 de '
        'Sentence Transformers para codificar consultas y entradas de la base de conocimiento.'
    )
    
    doc.add_heading('4.2 Generación de Embeddings', level=2)
    
    code = '''import torch
import torch.nn.functional as F
from transformers import AutoTokenizer, AutoModel

def encode_with_mean_pooling(model, tokenizer, texts):
    inputs = tokenizer(
        texts, padding="max_length", truncation=True,
        max_length=128, return_tensors="pt"
    )
    
    with torch.no_grad():
        outputs = model(**inputs)
        attention_mask = inputs['attention_mask']
        token_embeddings = outputs.last_hidden_state
        
        # Expandir máscara para broadcasting
        input_mask_expanded = attention_mask.unsqueeze(-1).expand(
            token_embeddings.size()
        ).float()
        
        # Mean pooling
        sum_embeddings = torch.sum(
            token_embeddings * input_mask_expanded, 1
        )
        sum_mask = torch.clamp(input_mask_expanded.sum(1), min=1e-9)
        embeddings = sum_embeddings / sum_mask
        
        # Normalización L2
        embeddings = F.normalize(embeddings, p=2, dim=1)
    
    return embeddings.numpy()'''
    add_code_block(doc, code, "Python")
    cap = doc.add_paragraph('Código 5: Mean pooling para embeddings de oraciones')
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    
    doc.add_heading('4.3 Pipeline de Procesamiento', level=2)
    doc.add_paragraph('El pipeline de búsqueda semántica procesa las consultas:')
    doc.add_paragraph('1. Tokenizar consulta de entrada (máx 128 tokens)', style='List Number')
    doc.add_paragraph('2. Codificar consulta usando codificador MiniLM', style='List Number')
    doc.add_paragraph('3. Calcular similitud coseno contra 517 embeddings pre-calculados', style='List Number')
    doc.add_paragraph('4. Recuperar top-3 contextos sobre umbral (0.4)', style='List Number')
    doc.add_paragraph('5. Aumentar prompt con contextos recuperados', style='List Number')
    doc.add_paragraph('6. Generar respuesta usando Llama 3.2', style='List Number')
    
    # Add RAG image
    img_path = os.path.join(os.path.dirname(__file__), 'fig_rag_pipeline.png')
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph('Figura 4: Pipeline de búsqueda semántica para generación de respuestas')
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].italic = True
    
    code = '''fun findTopKContexts(query: String, topK: Int = 3): List<Context> {
    val queryEmbedding = computeEmbedding(query)
    
    return kbEmbeddings.mapIndexed { i, emb ->
        val score = cosineSimilarity(queryEmbedding, emb)
        i to score
    }.filter { it.second >= 0.4f }
     .sortedByDescending { it.second }
     .take(topK)
     .map { (idx, score) ->
         Context(
             answer = kbEntries[idx].answer,
             score = score,
             category = kbEntries[idx].category
         )
     }
}

fun cosineSimilarity(a: FloatArray, b: FloatArray): Float {
    return a.indices.sumOf { (a[it] * b[it]).toDouble() }.toFloat()
}'''
    add_code_block(doc, code, "Kotlin")
    cap = doc.add_paragraph('Código 6: Implementación de búsqueda por similitud coseno')
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    
    # LLM
    doc.add_heading('5. LLM: Modelo de Lenguaje Local', level=1)
    
    doc.add_heading('5.1 Implementación del Modelo', level=2)
    doc.add_paragraph(
        'FarmifAI usa Llama 3.2 1B Instruct con cuantización Q4_K_M, reduciendo el tamaño del modelo '
        'de ~2GB a ~750MB manteniendo la calidad de respuestas.'
    )
    
    doc.add_heading('5.2 Inferencia', level=2)
    
    code = '''val prompt = """
<|begin_of_text|><|start_header_id|>system<|end_header_id|>

Eres FarmifAI, un asistente agrícola experto para 
agricultores colombianos. Responde de forma clara, 
práctica y concisa.

Contexto relevante:
$retrievedContext
<|eot_id|><|start_header_id|>user<|end_header_id|}

$userQuery<|eot_id|><|start_header_id|>assistant<|end_header_id|>
"""

val response = llamaService.generate(
    prompt = prompt,
    maxTokens = 256,
    temperature = 0.7f,
    topP = 0.9f
)'''
    add_code_block(doc, code, "Kotlin")
    cap = doc.add_paragraph('Código 7: Construcción de prompt y generación con Llama 3.2')
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    
    # Interfaz de Voz
    doc.add_heading('6. Interfaz de Voz', level=1)
    
    doc.add_heading('6.1 Reconocimiento de Voz', level=2)
    doc.add_paragraph(
        'La entrada de voz es manejada por Vosk, un toolkit de reconocimiento de voz offline liviano '
        'con un modelo de idioma español de 50MB.'
    )
    
    code = '''class VoiceHelper(context: Context) {
    private var voskModel: Model? = null
    private var speechService: SpeechService? = null
    
    fun initialize(modelPath: String) {
        voskModel = Model(modelPath)
    }
    
    fun startListening() {
        val recognizer = Recognizer(voskModel, 16000.0f)
        speechService = SpeechService(recognizer, 16000.0f)
        speechService?.startListening(object : RecognitionListener {
            override fun onResult(hypothesis: String?) {
                hypothesis?.let {
                    val json = JSONObject(it)
                    val text = json.getString("text")
                    if (text.isNotBlank()) onResult?.invoke(text)
                }
            }
        })
    }
}'''
    add_code_block(doc, code, "Kotlin")
    cap = doc.add_paragraph('Código 8: Integración de Vosk STT')
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    
    # Resultados
    doc.add_heading('7. Resultados y Evaluación', level=1)
    
    doc.add_heading('7.1 Métricas de Rendimiento', level=2)
    
    # Performance table
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'
    headers = ['Componente', 'Latencia', 'Memoria', 'Precisión']
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'DDDDDD')
    
    data = [
        ['Modelo de Visión', '200-400ms', '50MB', '92.3% Top-1'],
        ['Búsqueda Semántica', '60-130ms', '100MB', '94.4% Top-1'],
        ['Generación LLM', '2-5s', '1.2GB', '--'],
        ['Voz STT', 'Tiempo real', '100MB', '85-92%'],
    ]
    for row_idx, row_data in enumerate(data, 1):
        for col_idx, value in enumerate(row_data):
            table.cell(row_idx, col_idx).text = value
    
    cap = doc.add_paragraph()
    cap.add_run('Tabla 1: Métricas de rendimiento en dispositivo Android de gama media').italic = True
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('7.2 Resultados de Detección de Enfermedades', level=2)
    doc.add_paragraph('El modelo de visión logra:')
    doc.add_paragraph('• Precisión Top-1: 92.3%', style='List Bullet')
    doc.add_paragraph('• Precisión Top-3: 98.1%', style='List Bullet')
    doc.add_paragraph('• Tiempo de Inferencia: 200-400ms', style='List Bullet')
    
    # Despliegue
    doc.add_heading('8. Flujo de Trabajo de Despliegue', level=1)
    
    # Add deployment image
    img_path = os.path.join(os.path.dirname(__file__), 'fig_deployment_workflow.png')
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph('Figura 5: Flujo de trabajo de despliegue de modelos')
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].italic = True
    
    # Datasets
    doc.add_heading('9. Datasets y Fuentes de Datos', level=1)
    
    doc.add_heading('9.1 Datos de Entrenamiento del Modelo de Visión', level=2)
    doc.add_paragraph(
        'El modelo de clasificación de enfermedades fue entrenado usando dos datasets complementarios:'
    )
    
    p = doc.add_paragraph()
    p.add_run('Dataset PlantVillage: ').bold = True
    p.add_run('Un dataset públicamente disponible que contiene 54,309 imágenes de hojas de plantas sanas '
              'y enfermas en 14 especies de cultivos. Se extrajeron imágenes para cultivos objetivo (tomate, '
              'papa, pimiento, maíz) cubriendo 18 clases de enfermedades más estados saludables.')
    
    p = doc.add_paragraph()
    p.add_run('Disponible en: ').bold = True
    p.add_run('https://github.com/spMohanty/PlantVillage-Dataset y https://www.kaggle.com/datasets/emmarex/plantdisease')
    
    p = doc.add_paragraph()
    p.add_run('Dataset BRACOL de Café: ').bold = True
    p.add_run('El dataset Brazilian Arabica Coffee Leaf (BRACOL) contiene 1,747 imágenes de hojas de café '
              'arábica afectadas por los principales estreses bióticos, incluyendo cuatro clases de enfermedades: '
              'Minador de la hoja, Roya del café, Mancha parda y Mancha de Cercospora.')
    
    p = doc.add_paragraph()
    p.add_run('Disponible en: ').bold = True
    p.add_run('https://data.mendeley.com/datasets/yy2k5y8mxg/1 (Paper: https://doi.org/10.1016/j.compag.2019.105162)')
    
    doc.add_paragraph('Ambos datasets fueron preprocesados a 224×224 píxeles y aumentados usando rotaciones, '
                     'volteos, zooms y variaciones de brillo.')
    
    doc.add_heading('9.2 Base de Conocimiento NLP', level=2)
    doc.add_paragraph('El componente de búsqueda semántica usa una base de conocimiento agrícola curada:')
    doc.add_paragraph('• 517 pares pregunta-respuesta cubriendo manejo de cultivos, identificación de enfermedades, control de plagas', style='List Bullet')
    doc.add_paragraph('• Preguntas obtenidas de servicios de extensión agrícola, foros de agricultores y consultas con expertos', style='List Bullet')
    doc.add_paragraph('• Contenido validado por agrónomos de la Universidad del Cauca', style='List Bullet')
    doc.add_paragraph('• Categorías: Manejo de Cultivos (35%), Control de Enfermedades (30%), Manejo de Plagas (20%), Suelo y Fertilización (15%)', style='List Bullet')
    
    doc.add_heading('9.3 Modelo de Reconocimiento de Voz', level=2)
    doc.add_paragraph('El reconocimiento de voz offline usa el modelo de idioma español pre-entrenado de Vosk (vosk-model-small-es-0.42):')
    doc.add_paragraph('• Datos de entrenamiento: Corpus Common Voice en español + datasets propietarios', style='List Bullet')
    doc.add_paragraph('• Tamaño del modelo: 50MB comprimido', style='List Bullet')
    doc.add_paragraph('• Optimizado para dialectos del español latinoamericano', style='List Bullet')
    doc.add_paragraph('• Vocabulario adaptado incluye terminología agrícola', style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('Disponible en: ').bold = True
    p.add_run('https://alphacephei.com/vosk/models/vosk-model-small-es-0.42.zip (Modelos oficiales: https://alphacephei.com/vosk/models)')
    
    doc.add_heading('9.4 Modelo de Lenguaje', level=2)
    doc.add_paragraph('El modelo Llama 3.2 1B Instruct fue desarrollado por Meta AI:')
    doc.add_paragraph('• Pre-entrenado en texto diverso de internet y código', style='List Bullet')
    doc.add_paragraph('• Ajustado para tareas de seguimiento de instrucciones', style='List Bullet')
    doc.add_paragraph('• Cuantizado a formato Q4_K_M usando llama.cpp', style='List Bullet')
    doc.add_paragraph('• No se aplicó ajuste fino adicional específico del dominio', style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('Disponible en: ').bold = True
    p.add_run('https://huggingface.co/meta-llama/Llama-3.2-1B-Instruct (Cuantizado: https://huggingface.co/bartowski/Llama-3.2-1B-Instruct-GGUF)')
    
    # Save
    output_path = os.path.join(os.path.dirname(__file__), 'FarmifAI_Documentation_ES.docx')
    doc.save(output_path)
    print(f"Spanish document saved: {output_path}")
    return output_path


if __name__ == '__main__':
    print("Generating Word documents...")
    en_path = create_english_doc()
    es_path = create_spanish_doc()
    print(f"\nDocuments created:")
    print(f"  English: {en_path}")
    print(f"  Spanish: {es_path}")
